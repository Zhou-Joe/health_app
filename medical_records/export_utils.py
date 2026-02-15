"""导出工具 - 用于生成PDF和Word文档"""
import os
import re
from datetime import datetime
from django.conf import settings
from django.http import HttpResponse
from urllib.parse import quote
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .models import Conversation, HealthAdvice, HealthIndicator, HealthCheckup


def register_chinese_font():
    """注册中文字体"""
    font_registered = False

    # 强制使用 wqy-microhei 字体（优先检查最常见的安装路径）
    font_paths = [
        # RHEL/CentOS/Alibaba Cloud Linux - wqy-microhei
        '/usr/share/fonts/wqy-microhei/wqy-microhei.ttc',  # 主要路径
        '/usr/share/fonts/wqy-microhei/wqy-microhei.tbf',  # 备选扩展名
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',  # Ubuntu/Debian 路径
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttf',  # Ubuntu/Debian 路径
        '/usr/share/fonts/wqy/wqy-microhei.ttc',  # 简化路径

        # 如果上面的都不行，尝试其他字体
        '/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc',  # Noto Sans CJK
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # WQY Zenhei
        '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',  # Droid
    ]

    for font_path in font_paths:
        try:
            if os.path.exists(font_path):
                # 注册字体，使用别名
                pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                font_registered = True
                print(f"✓ Successfully registered Chinese font: {font_path}")
                break
        except Exception as e:
            print(f"✗ Failed to register font {font_path}: {str(e)}")
            continue

    if not font_registered:
        print("⚠ Warning: No Chinese font found. Chinese characters may display as squares.")
        print("Please install wqy-microhei-fonts:")
        print("  - RHEL/CentOS/Alibaba Cloud: sudo yum install -y wqy-microhei-fonts")

    return font_registered


# 尝试注册中文字体
CHINESE_FONT_AVAILABLE = register_chinese_font()


def markdown_to_pdf_text(text):
    """
    将 Markdown 格式转换为 PDF 可显示的格式（使用 HTML 标签）
    """
    # 转义 HTML 特殊字符
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')

    # 替换 Markdown 格式为 HTML 标签
    # 粗体 **text** 或 __text__
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.+?)__', r'<b>\1</b>', text)

    # 斜体 *text* 或 _text_
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)
    text = re.sub(r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'<i>\1</i>', text)

    # 粗斜体 ***text*** 或 ___text___
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    text = re.sub(r'___(.+?)___', r'<b><i>\1</i></b>', text)

    # 删除线 ~~text~~
    text = re.sub(r'~~(.+?)~~', r'<u><strike>\1</strike></u>', text)

    # 代码 `text`
    text = re.sub(r'`(.+?)`', r'<font face="Courier" color="#e83e8c"><i>\1</i></font>', text)

    # 标题 # 标题
    text = re.sub(r'^### (.+)$', r'<b><font size="12">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'<b><font size="14">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.+)$', r'<b><font size="16">\1</font></b>', text, flags=re.MULTILINE)

    # 处理换行
    text = text.replace('\n', '<br/>')

    return text



def add_formatted_text_to_paragraph(paragraph, text):
    """
    将 Markdown 格式的文本添加到段落，保留格式（加粗、斜体、代码等）
    """
    # 定义 Markdown 正则模式
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # ***粗斜体***
        (r'___(.+?)___', 'bold_italic'),         # ___粗斜体___
        (r'\*\*(.+?)\*\*', 'bold'),              # **粗体**
        (r'__(.+?)__', 'bold'),                  # __粗体__
        (r'\*(.+?)\*', 'italic'),                # *斜体*
        (r'_(.+?)_', 'italic'),                  # _斜体_
        (r'~~(.+?)~~', 'strike'),                # ~~删除线~~
        (r'`(.+?)`', 'code'),                    # `代码`
    ]

    # 按照在文本中出现的位置排序所有匹配
    matches = []
    for pattern, style_type in patterns:
        for match in re.finditer(pattern, text, re.DOTALL):
            matches.append({
                'start': match.start(),
                'end': match.end(),
                'content': match.group(1),
                'type': style_type,
                'full_match': match.group(0)
            })

    # 如果没有匹配到任何格式，直接添加纯文本
    if not matches:
        paragraph.add_run(text)
        return

    # 按起始位置排序
    matches.sort(key=lambda x: x['start'])

    # 合并重叠的匹配
    merged_matches = []
    for match in matches:
        if merged_matches and match['start'] < merged_matches[-1]['end']:
            # 有重叠，跳过这个匹配（保留前一个）
            continue
        merged_matches.append(match)

    # 构建带格式的文本
    last_end = 0
    for match in merged_matches:
        # 添加前面的普通文本
        if match['start'] > last_end:
            paragraph.add_run(text[last_end:match['start']])

        # 添加带格式的文本
        run = paragraph.add_run(match['content'])
        if match['type'] == 'bold':
            run.bold = True
        elif match['type'] == 'italic':
            run.italic = True
        elif match['type'] == 'bold_italic':
            run.bold = True
            run.italic = True
        elif match['type'] == 'strike':
            run.font.strike = True
        elif match['type'] == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(232, 62, 140)
            # 添加浅灰色背景
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F1F2F6')
            run._element.get_or_add_rPr().append(shading_elm)

        last_end = match['end']

    # 添加剩余的普通文本
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def add_markdown_paragraphs(doc, markdown_text):
    """
    将 Markdown 文本添加到 Word 文档，保留段落和格式
    """
    # 按行分割
    lines = markdown_text.split('\n')
    current_paragraph = None
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 空行
        if not line:
            if current_paragraph:
                current_paragraph = None
            doc.add_paragraph()  # 添加空行
            i += 1
            continue

        # 标题行（# 标记）
        if line.startswith('#'):
            level = min(len(re.match(r'^#+', line).group()), 6)
            title_text = line.lstrip('#').strip()
            # 标题只添加纯文本，不支持复杂格式（避免 Run 对象问题）
            heading = doc.add_heading(title_text, level=level)
            current_paragraph = None
            i += 1
            continue

        # 列表项（- 或 * 或 数字.）
        if re.match(r'^\s*[-*]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'^\s*[-*\d.]+\s+', '', line)
            add_formatted_text_to_paragraph(p, text)
            current_paragraph = None
            i += 1
            continue

        # 代码块（``` 包围）
        if line.startswith('```'):
            # 收集代码块内容
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1

            p = doc.add_paragraph('\n'.join(code_lines))
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(226, 232, 240)
            # 设置段落背景（深色）
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '2D3748')
            p._element.get_or_add_pPr().append(shading_elm)
            current_paragraph = None
            i += 1  # 跳过结束的 ```
            continue

        # 普通段落
        if current_paragraph is None:
            current_paragraph = doc.add_paragraph()
        else:
            # 同一段落内的新行
            current_paragraph.add_run('\n')

        add_formatted_text_to_paragraph(current_paragraph, line)
        i += 1


class ConversationExporter:
    """对话导出器"""

    def __init__(self, conversation_id):
        self.conversation = Conversation.objects.get(id=conversation_id)
        self.messages = HealthAdvice.get_conversation_messages(conversation_id)
        self.title = self.conversation.title

    def export_to_pdf(self):
        """导出为PDF"""
        response = HttpResponse(content_type='application/pdf')
        filename = f"AI健康咨询_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.pdf"
        # 对中文文件名进行编码，兼容不同浏览器
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        # 创建PDF文档
        doc = SimpleDocTemplate(response, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()

        # 创建样式（使用中文字体）
        if CHINESE_FONT_AVAILABLE:
            # 有中文字体时使用中文字体
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Heading1'],
                fontName='ChineseFont',
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
                leading=22
            )
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=11,
                leading=16,
                alignment=TA_LEFT
            )
            question_style = ParagraphStyle(
                'Question',
                parent=normal_style,
                fontName='ChineseFont',
                fontSize=12,
                textColor=colors.HexColor('#0066cc'),
                spaceBefore=10,
                spaceAfter=5
            )
            answer_style = ParagraphStyle(
                'Answer',
                parent=normal_style,
                fontName='ChineseFont',
                fontSize=12,
                textColor=colors.HexColor('#009900'),
                spaceBefore=5,
                spaceAfter=5
            )
            label_font = 'ChineseFont'
        else:
            # 无中文字体时使用默认字体（中文可能显示为方块）
            title_style = styles['Heading1']
            normal_style = styles['Normal']
            question_style = ParagraphStyle(
                'Question',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#0066cc'),
                spaceBefore=10,
                spaceAfter=5
            )
            answer_style = ParagraphStyle(
                'Answer',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#009900'),
                spaceBefore=5,
                spaceAfter=5
            )
            label_font = 'Helvetica-Bold'

        # 标题
        story.append(Paragraph(self.title, title_style))
        story.append(Spacer(1, 0.5 * cm))

        # 对话信息表格
        info_data = [
            ['创建时间', self.conversation.created_at.strftime('%Y年%m月%d日 %H:%M:%S')],
            ['更新时间', self.conversation.updated_at.strftime('%Y年%m月%d日 %H:%M:%S')],
            ['消息数量', str(len(self.messages)) + ' 条'],
        ]

        info_table = Table(info_data, colWidths=[4*cm, 11*cm])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), label_font if CHINESE_FONT_AVAILABLE else 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 1 * cm))

        # 对话内容
        for idx, msg in enumerate(self.messages, 1):
            # 问题标题
            story.append(Paragraph(f"问题 {idx}:", question_style))

            # 问题内容（支持 Markdown 格式）
            question_html = markdown_to_pdf_text(msg.question)
            story.append(Paragraph(question_html, normal_style))
            story.append(Spacer(1, 0.3 * cm))

            # 回答标题
            story.append(Paragraph(f"回答 {idx}:", answer_style))

            # 回答内容（支持 Markdown 格式）
            answer_html = markdown_to_pdf_text(msg.answer)
            story.append(Paragraph(answer_html, normal_style))
            story.append(Spacer(1, 0.8 * cm))

        # 构建PDF
        doc.build(story)
        return response

    def export_to_word(self):
        """导出为Word"""
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"AI健康咨询_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.docx"
        # 对中文文件名进行编码，兼容不同浏览器
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        # 创建Word文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        # 标题
        title = doc.add_heading(self.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 26, 26)

        # 对话信息
        doc.add_paragraph('对话信息：')
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ['创建时间', self.conversation.created_at.strftime('%Y-%m-%d %H:%M:%S')],
            ['更新时间', self.conversation.updated_at.strftime('%Y-%m-%d %H:%M:%S')],
            ['消息数量', str(len(self.messages))],
        ]

        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # 设置第一列为粗体
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        # 对话内容
        for idx, msg in enumerate(self.messages, 1):
            # 问题
            q_heading = doc.add_heading(f'问题 {idx}:', level=3)
            q_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

            # 使用 Markdown 解析添加问题
            q_para = doc.add_paragraph()
            q_para.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text_to_paragraph(q_para, msg.question)

            # 回答
            a_heading = doc.add_heading(f'回答 {idx}:', level=3)
            a_heading.runs[0].font.color.rgb = RGBColor(0, 153, 0)

            # 使用 Markdown 解析添加回答（支持多段落和格式）
            answer_container = doc.add_paragraph()
            answer_container.paragraph_format.left_indent = Inches(0.25)

            # 创建一个临时文档来解析 Markdown，然后将内容复制过来
            temp_doc = Document()
            add_markdown_paragraphs(temp_doc, msg.answer)

            # 将临时文档的段落内容复制到主文档
            for temp_para in temp_doc.paragraphs:
                if not temp_para.text.strip() and len(temp_para.runs) == 0:
                    # 空段落
                    new_para = doc.add_paragraph()
                    new_para.paragraph_format.left_indent = Inches(0.25)
                else:
                    # 有内容的段落
                    new_para = doc.add_paragraph()
                    new_para.paragraph_format.left_indent = Inches(0.25)

                    # 复制所有 runs
                    for run in temp_para.runs:
                        new_run = new_para.add_run(run.text)
                        # 复制格式
                        if run.bold:
                            new_run.bold = True
                        if run.italic:
                            new_run.italic = True
                        if run.font.strike:
                            new_run.font.strike = True
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb

                    # 复制段落样式
                    if temp_para.style and temp_para.style.name != 'Normal':
                        try:
                            new_para.style = temp_para.style
                        except:
                            pass

            # 添加分隔线
            if idx < len(self.messages):
                sep_para = doc.add_paragraph('_' * 80)
                sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到响应
        doc.save(response)
        return response


class AISummaryExporter:
    """AI对话总结导出器"""

    def __init__(self, conversation_id):
        self.conversation = Conversation.objects.get(id=conversation_id)
        self.messages = HealthAdvice.get_conversation_messages(conversation_id)
        self.title = self.conversation.title
        self.ai_summary = self.conversation.ai_summary
        self.ai_summary_created_at = self.conversation.ai_summary_created_at

    def export_to_pdf(self):
        """导出AI总结为PDF"""
        response = HttpResponse(content_type='application/pdf')
        filename = f"AI对话总结_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.pdf"
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        doc = SimpleDocTemplate(response, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()

        if CHINESE_FONT_AVAILABLE:
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Heading1'],
                fontName='ChineseFont',
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
                leading=22
            )
            heading_style = ParagraphStyle(
                'ChineseHeading',
                parent=styles['Heading2'],
                fontName='ChineseFont',
                fontSize=14,
                textColor=colors.HexColor('#0066cc'),
                spaceAfter=8,
                leading=18
            )
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=11,
                leading=16,
                alignment=TA_LEFT
            )
            label_font = 'ChineseFont'
        else:
            title_style = styles['Heading1']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            label_font = 'Helvetica-Bold'

        story.append(Paragraph(f"{self.title} - AI总结", title_style))
        story.append(Spacer(1, 0.5 * cm))

        info_data = [
            ['对话创建时间', self.conversation.created_at.strftime('%Y年%m月%d日 %H:%M')],
            ['对话更新时间', self.conversation.updated_at.strftime('%Y年%m月%d日 %H:%M')],
            ['消息数量', f'{len(self.messages)} 条'],
        ]
        if self.ai_summary_created_at:
            info_data.append(['总结生成时间', self.ai_summary_created_at.strftime('%Y年%m月%d日 %H:%M')])

        info_table = Table(info_data, colWidths=[4*cm, 11*cm])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), label_font if CHINESE_FONT_AVAILABLE else 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 1 * cm))

        if self.ai_summary:
            story.append(Paragraph('AI总结内容', heading_style))
            story.append(Spacer(1, 0.3 * cm))
            summary_html = markdown_to_pdf_text(self.ai_summary)
            story.append(Paragraph(summary_html, normal_style))
        else:
            story.append(Paragraph('暂无AI总结内容', normal_style))

        doc.build(story)
        return response

    def export_to_word(self):
        """导出AI总结为Word"""
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"AI对话总结_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.docx"
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        title = doc.add_heading(f"{self.title} - AI总结", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 26, 26)

        doc.add_paragraph('对话信息：')
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ['对话创建时间', self.conversation.created_at.strftime('%Y-%m-%d %H:%M')],
            ['对话更新时间', self.conversation.updated_at.strftime('%Y-%m-%d %H:%M')],
            ['消息数量', str(len(self.messages))],
        ]
        if self.ai_summary_created_at:
            info_data.append(['总结生成时间', self.ai_summary_created_at.strftime('%Y-%m-%d %H:%M')])
        else:
            info_data.append(['总结生成时间', '暂无'])

        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        if self.ai_summary:
            doc.add_heading('AI总结内容', level=2)
            summary_para = doc.add_paragraph()
            add_formatted_text_to_paragraph(summary_para, self.ai_summary)
        else:
            doc.add_paragraph('暂无AI总结内容')

        doc.save(response)
        return response


class EventAiSummaryExporter:
    """健康事件AI分析导出器"""

    def __init__(self, event_id):
        self.event = HealthEvent.objects.get(id=event_id)
        self.ai_summary = self.event.ai_summary
        self.ai_summary_created_at = self.event.ai_summary_created_at

    def export_to_pdf(self):
        """导出事件AI分析为PDF"""
        response = HttpResponse(content_type='application/pdf')
        filename = f"健康事件分析_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.pdf"
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        doc = SimpleDocTemplate(response, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()

        if CHINESE_FONT_AVAILABLE:
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Heading1'],
                fontName='ChineseFont',
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
                leading=22
            )
            heading_style = ParagraphStyle(
                'ChineseHeading',
                parent=styles['Heading2'],
                fontName='ChineseFont',
                fontSize=14,
                textColor=colors.HexColor('#0066cc'),
                spaceAfter=8,
                leading=18
            )
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=11,
                leading=16,
                alignment=TA_LEFT
            )
            label_font = 'ChineseFont'
        else:
            title_style = styles['Heading1']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            label_font = 'Helvetica-Bold'

        story.append(Paragraph(f"{self.event.name} - AI健康分析", title_style))
        story.append(Spacer(1, 0.5 * cm))

        event_type_display = dict(HealthEvent.EVENT_TYPE_CHOICES).get(self.event.event_type, self.event.event_type)
        status_display = dict(HealthEvent.EVENT_STATUS_CHOICES).get(self.event.status, self.event.status)

        info_data = [
            ['事件名称', self.event.name],
            ['事件类型', event_type_display],
            ['开始日期', self.event.start_date.strftime('%Y年%m月%d日')],
            ['结束日期', self.event.end_date.strftime('%Y年%m月%d日') if self.event.end_date else '进行中'],
            ['当前状态', status_display],
        ]
        if self.ai_summary_created_at:
            info_data.append(['分析生成时间', self.ai_summary_created_at.strftime('%Y年%m月%d日 %H:%M')])

        info_table = Table(info_data, colWidths=[4*cm, 11*cm])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), label_font if CHINESE_FONT_AVAILABLE else 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 1 * cm))

        if self.event.description:
            story.append(Paragraph('事件描述', heading_style))
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph(self.event.description, normal_style))
            story.append(Spacer(1, 0.5 * cm))

        if self.ai_summary:
            story.append(Paragraph('AI健康分析', heading_style))
            story.append(Spacer(1, 0.3 * cm))
            summary_html = markdown_to_pdf_text(self.ai_summary)
            story.append(Paragraph(summary_html, normal_style))
        else:
            story.append(Paragraph('暂无AI分析内容', normal_style))

        doc.build(story)
        return response

    def export_to_word(self):
        """导出事件AI分析为Word"""
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"健康事件分析_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.docx"
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        title = doc.add_heading(f"{self.event.name} - AI健康分析", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 26, 26)

        event_type_display = dict(HealthEvent.EVENT_TYPE_CHOICES).get(self.event.event_type, self.event.event_type)
        status_display = dict(HealthEvent.EVENT_STATUS_CHOICES).get(self.event.status, self.event.status)

        doc.add_paragraph('事件信息：')
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ['事件名称', self.event.name],
            ['事件类型', event_type_display],
            ['开始日期', self.event.start_date.strftime('%Y-%m-%d')],
            ['结束日期', self.event.end_date.strftime('%Y-%m-%d') if self.event.end_date else '进行中'],
            ['当前状态', status_display],
        ]
        if self.ai_summary_created_at:
            info_data.append(['分析生成时间', self.ai_summary_created_at.strftime('%Y-%m-%d %H:%M')])
        else:
            info_data.append(['分析生成时间', '暂无'])

        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        if self.event.description:
            doc.add_heading('事件描述', level=2)
            doc.add_paragraph(self.event.description)
            doc.add_paragraph()

        if self.ai_summary:
            doc.add_heading('AI健康分析', level=2)
            summary_para = doc.add_paragraph()
            add_formatted_text_to_paragraph(summary_para, self.ai_summary)
        else:
            doc.add_paragraph('暂无AI分析内容')

        doc.save(response)
        return response


class HealthTrendsExporter:
    """健康趋势数据导出器"""

    def __init__(self, user):
        self.user = user
        self.indicators = self._get_indicators_by_type()

    def _get_indicators_by_type(self):
        """按类型分组获取健康指标"""
        from .models import HealthIndicator

        all_indicators = HealthIndicator.objects.filter(
            checkup__user=self.user
        ).select_related('checkup').order_by('checkup__checkup_date', 'indicator_name')

        # 按类型分组
        grouped = {}
        for indicator in all_indicators:
            if indicator.indicator_type not in grouped:
                grouped[indicator.indicator_type] = []
            grouped[indicator.indicator_type].append(indicator)

        return grouped

    def export_to_pdf(self):
        """导出为PDF"""
        response = HttpResponse(content_type='application/pdf')
        filename = f"健康趋势分析_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.pdf"
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        # 创建PDF文档
        doc = SimpleDocTemplate(response, pagesize=A4,
                               leftMargin=1.5*cm, rightMargin=1.5*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()

        # 样式设置
        if CHINESE_FONT_AVAILABLE:
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontName='ChineseFont',
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
                leading=22,
                alignment=TA_CENTER
            )
            category_style = ParagraphStyle(
                'Category',
                parent=styles['Heading2'],
                fontName='ChineseFont',
                fontSize=14,
                textColor=colors.HexColor('#0066cc'),
                spaceBefore=15,
                spaceAfter=10,
                leading=18
            )
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=10,
                leading=14
            )
            table_font = 'ChineseFont'
        else:
            title_style = styles['Heading1']
            category_style = styles['Heading2']
            normal_style = styles['Normal']
            table_font = 'Helvetica'

        # 标题
        story.append(Paragraph("健康趋势分析报告", title_style))
        story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}", normal_style))
        story.append(Spacer(1, 0.5 * cm))

        # 按类型导出数据
        type_names = {
            'physical_exam': '体格检查',
            'blood_routine': '血液常规',
            'biochemistry': '生化检验',
            'liver_function': '肝功能',
            'kidney_function': '肾功能',
            'thyroid_function': '甲状腺功能',
            'tumor_markers': '肿瘤标志物',
            'urine_exam': '尿液检查',
            'blood_rheology': '血液流变学',
            'eye_exam': '眼科检查',
            'ultrasound_exam': '超声检查',
            'imaging_exam': '影像学检查',
            'diagnosis': '病症诊断',
            'symptoms': '症状描述',
            'other_exam': '其他检查',
        }

        first_category = True
        for indicator_type, indicators in self.indicators.items():
            if not indicators:
                continue

            if not first_category:
                story.append(PageBreak())
            first_category = False

            # 分类标题
            type_name = type_names.get(indicator_type, indicator_type)
            story.append(Paragraph(type_name, category_style))

            # 按指标名称分组
            indicator_groups = {}
            for ind in indicators:
                if ind.indicator_name not in indicator_groups:
                    indicator_groups[ind.indicator_name] = []
                indicator_groups[ind.indicator_name].append(ind)

            # 为每个指标创建表格
            for indicator_name, ind_list in indicator_groups.items():
                # 表格标题
                story.append(Paragraph(f"【{indicator_name}】", ParagraphStyle(
                    'IndicatorTitle',
                    parent=category_style,
                    fontName=table_font if CHINESE_FONT_AVAILABLE else 'Helvetica-Bold',
                    fontSize=12,
                    textColor=colors.HexColor('#333333'),
                    spaceBefore=10,
                    spaceAfter=5
                )))

                # 准备表格数据
                table_data = [['体检日期', '检测值', '单位', '参考范围', '状态']]
                for ind in ind_list:
                    date_str = ind.checkup.checkup_date.strftime('%Y-%m-%d')
                    value = ind.value or '-'
                    unit = ind.unit or '-'
                    ref_range = ind.reference_range or '-'
                    status_map = {'normal': '正常', 'abnormal': '异常', 'attention': '关注'}
                    status = status_map.get(ind.status, ind.status)

                    table_data.append([date_str, value, unit, ref_range, status])

                # 创建表格
                table = Table(table_data, colWidths=[2.5*cm, 2*cm, 2*cm, 4*cm, 1.5*cm])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('FONTNAME', (0, 0), (-1, -1), table_font if CHINESE_FONT_AVAILABLE else 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                ]))

                # 状态列着色
                for i, row in enumerate(table_data[1:], start=1):
                    status = row[4]
                    if status == '异常':
                        table.setStyle(TableStyle([('TEXTCOLOR', (0, i), (-1, i), colors.red)]))
                    elif status == '关注':
                        table.setStyle(TableStyle([('TEXTCOLOR', (0, i), (-1, i), colors.HexColor('#ff9800'))]))

                story.append(table)
                story.append(Spacer(1, 0.3 * cm))

        # 构建PDF
        doc.build(story)
        return response

    def export_to_word(self):
        """导出为Word"""
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"健康趋势分析_{datetime.now().strftime('%Y年%m月%d日_%H%M')}.docx"
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        # 创建Word文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        # 标题
        title = doc.add_heading('健康趋势分析报告', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 26, 26)

        # 生成时间
        doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        doc.add_paragraph()

        # 类型名称映射
        type_names = {
            'physical_exam': '体格检查',
            'blood_routine': '血液常规',
            'biochemistry': '生化检验',
            'liver_function': '肝功能',
            'kidney_function': '肾功能',
            'thyroid_function': '甲状腺功能',
            'tumor_markers': '肿瘤标志物',
            'urine_exam': '尿液检查',
            'blood_rheology': '血液流变学',
            'eye_exam': '眼科检查',
            'ultrasound_exam': '超声检查',
            'imaging_exam': '影像学检查',
            'diagnosis': '病症诊断',
            'symptoms': '症状描述',
            'other_exam': '其他检查',
        }

        status_colors = {
            'normal': RGBColor(40, 167, 69),    # 绿色
            'abnormal': RGBColor(220, 53, 69),  # 红色
            'attention': RGBColor(255, 152, 0)  # 橙色
        }

        # 按类型导出数据
        for indicator_type, indicators in self.indicators.items():
            if not indicators:
                continue

            # 分类标题
            type_name = type_names.get(indicator_type, indicator_type)
            category_heading = doc.add_heading(type_name, level=2)
            category_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

            # 按指标名称分组
            indicator_groups = {}
            for ind in indicators:
                if ind.indicator_name not in indicator_groups:
                    indicator_groups[ind.indicator_name] = []
                indicator_groups[ind.indicator_name].append(ind)

            # 为每个指标创建表格
            for indicator_name, ind_list in indicator_groups.items():
                # 指标名称
                ind_heading = doc.add_heading(f'【{indicator_name}】', level=3)

                # 准备表格数据
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'

                # 表头
                header_cells = table.rows[0].cells
                headers = ['体检日期', '检测值', '单位', '参考范围', '状态']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True

                # 数据行
                for ind in ind_list:
                    row_cells = table.add_row().cells
                    date_str = ind.checkup.checkup_date.strftime('%Y年%m月%d日')
                    value = ind.value or '-'
                    unit = ind.unit or '-'
                    ref_range = ind.reference_range or '-'
                    status_map = {'normal': '正常', 'abnormal': '异常', 'attention': '关注'}
                    status = status_map.get(ind.status, ind.status)

                    row_cells[0].text = date_str
                    row_cells[1].text = str(value)
                    row_cells[2].text = unit
                    row_cells[3].text = ref_range
                    row_cells[4].text = status

                    # 设置状态颜色
                    if ind.status in status_colors:
                        row_cells[4].paragraphs[0].runs[0].font.color.rgb = status_colors[ind.status]

                doc.add_paragraph()  # 空行

        # 保存到响应
        doc.save(response)
        return response


class CheckupReportsExporter:
    """批量导出体检报告"""

    def __init__(self, checkups):
        """
        初始化导出器
        :param checkups: HealthCheckup QuerySet 或列表
        """
        self.checkups = checkups
        self.user = checkups[0].user if checkups else None

        # 收集所有体检报告的指标数据
        self.indicators = {}
        for checkup in checkups:
            for indicator in checkup.indicators.all().order_by('indicator_type', 'indicator_name'):
                if indicator.indicator_type not in self.indicators:
                    self.indicators[indicator.indicator_type] = []
                self.indicators[indicator.indicator_type].append(indicator)

    def export_to_pdf(self):
        """导出为PDF"""
        response = HttpResponse(content_type='application/pdf')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'体检报告汇总_{timestamp}.pdf'
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        # 创建PDF文档
        doc = SimpleDocTemplate(response, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        # 自定义样式
        if CHINESE_FONT_AVAILABLE:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='ChineseFont',
                fontSize=20,
                textColor=colors.HexColor('#0066cc'),
                spaceAfter=20,
                alignment=TA_CENTER
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName='ChineseFont',
                fontSize=14,
                textColor=colors.HexColor('#0066cc'),
                spaceAfter=10,
                spaceBefore=15
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=10,
                spaceAfter=5
            )
        else:
            title_style = styles['Heading1']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']

        # 标题
        title = Paragraph("体检报告汇总", title_style)
        story.append(title)
        story.append(Spacer(1, 0.5*cm))

        # 汇总信息
        if self.user:
            info_text = f"用户：{self.user.username} | 共 {len(self.checkups)} 份报告"
            story.append(Paragraph(info_text, normal_style))
            story.append(Spacer(1, 0.5*cm))

        # 类型名称映射
        type_names = {
            'general_exam': '一般检查',
            'blood_routine': '血常规',
            'biochemistry': '生化检验',
            'liver_function': '肝功能',
            'kidney_function': '肾功能',
            'thyroid': '甲状腺',
            'cardiac': '心脏标志物',
            'tumor_markers': '肿瘤标志物',
            'infection': '感染炎症',
            'blood_rheology': '血液流变',
            'coagulation': '凝血功能',
            'urine': '尿液检查',
            'stool': '粪便检查',
            'pathology': '病理检查',
            'ultrasound': '超声检查',
            'X_ray': 'X线检查',
            'CT_MRI': 'CT和MRI',
            'endoscopy': '内镜检查',
            'special_organs': '专科检查',
            'other': '其他检查',
        }

        # 按类型导出数据
        for indicator_type, indicators in self.indicators.items():
            if not indicators:
                continue

            # 分类标题
            type_name = type_names.get(indicator_type, indicator_type)
            story.append(Paragraph(type_name, heading_style))

            # 按指标名称分组
            indicator_groups = {}
            for ind in indicators:
                if ind.indicator_name not in indicator_groups:
                    indicator_groups[ind.indicator_name] = []
                indicator_groups[ind.indicator_name].append(ind)

            # 为每个指标创建表格
            for indicator_name, ind_list in indicator_groups.items():
                # 指标名称
                story.append(Paragraph(f"<b>{indicator_name}</b>", normal_style))

                # 准备表格数据
                table_data = [['体检日期', '检测值', '单位', '参考范围', '状态']]

                for ind in ind_list:
                    date_str = ind.checkup.checkup_date.strftime('%Y-%m-%d')
                    value = str(ind.value) if ind.value else '-'
                    unit = ind.unit or '-'
                    ref_range = ind.reference_range or '-'
                    status_map = {'normal': '正常', 'abnormal': '异常', 'attention': '关注'}
                    status = status_map.get(ind.status, ind.status)

                    table_data.append([date_str, value, unit, ref_range, status])

                # 创建表格
                table = Table(table_data, colWidths=[3*cm, 2.5*cm, 2*cm, 3*cm, 2*cm])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'ChineseFont' if CHINESE_FONT_AVAILABLE else 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'ChineseFont' if CHINESE_FONT_AVAILABLE else 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
                ]))

                story.append(table)
                story.append(Spacer(1, 0.3*cm))

            story.append(PageBreak())

        # 生成PDF
        doc.build(story)
        return response

    def export_to_word(self):
        """导出为Word"""
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'体检报告汇总_{timestamp}.docx'
        encoded_filename = quote(filename)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"

        # 创建Word文档
        doc = Document()

        # 设置默认字体（中文支持）
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 标题
        title = doc.add_heading('体检报告汇总', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 汇总信息
        if self.user:
            info_para = doc.add_paragraph()
            info_para.add_run(f"用户：{self.user.username} | 共 {len(self.checkups)} 份报告").font.size = Pt(11)

        doc.add_paragraph()  # 空行

        # 类型名称映射
        type_names = {
            'general_exam': '一般检查',
            'blood_routine': '血常规',
            'biochemistry': '生化检验',
            'liver_function': '肝功能',
            'kidney_function': '肾功能',
            'thyroid': '甲状腺',
            'cardiac': '心脏标志物',
            'tumor_markers': '肿瘤标志物',
            'infection': '感染炎症',
            'blood_rheology': '血液流变',
            'coagulation': '凝血功能',
            'urine': '尿液检查',
            'stool': '粪便检查',
            'pathology': '病理检查',
            'ultrasound': '超声检查',
            'X_ray': 'X线检查',
            'CT_MRI': 'CT和MRI',
            'endoscopy': '内镜检查',
            'special_organs': '专科检查',
            'other': '其他检查',
        }

        status_colors = {
            'normal': RGBColor(40, 167, 69),    # 绿色
            'abnormal': RGBColor(220, 53, 69),  # 红色
            'attention': RGBColor(255, 152, 0)  # 橙色
        }

        # 按类型导出数据
        for indicator_type, indicators in self.indicators.items():
            if not indicators:
                continue

            # 分类标题
            type_name = type_names.get(indicator_type, indicator_type)
            category_heading = doc.add_heading(type_name, level=2)
            category_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

            # 按指标名称分组
            indicator_groups = {}
            for ind in indicators:
                if ind.indicator_name not in indicator_groups:
                    indicator_groups[ind.indicator_name] = []
                indicator_groups[ind.indicator_name].append(ind)

            # 为每个指标创建表格
            for indicator_name, ind_list in indicator_groups.items():
                # 指标名称
                ind_heading = doc.add_heading(f'【{indicator_name}】', level=3)

                # 准备表格数据
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'

                # 表头
                header_cells = table.rows[0].cells
                headers = ['体检日期', '检测值', '单位', '参考范围', '状态']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True

                # 数据行
                for ind in ind_list:
                    row_cells = table.add_row().cells
                    date_str = ind.checkup.checkup_date.strftime('%Y年%m月%d日')
                    value = ind.value or '-'
                    unit = ind.unit or '-'
                    ref_range = ind.reference_range or '-'
                    status_map = {'normal': '正常', 'abnormal': '异常', 'attention': '关注'}
                    status = status_map.get(ind.status, ind.status)

                    row_cells[0].text = date_str
                    row_cells[1].text = str(value)
                    row_cells[2].text = unit
                    row_cells[3].text = ref_range
                    row_cells[4].text = status

                    # 设置状态颜色
                    if ind.status in status_colors:
                        row_cells[4].paragraphs[0].runs[0].font.color.rgb = status_colors[ind.status]

                doc.add_paragraph()  # 空行

        # 保存到响应
        doc.save(response)
        return response
